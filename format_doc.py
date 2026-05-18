import docx

doc_path = r"C:\Users\leoga\Desktop\Doctorado 2025\Artículos\Articulos Finales Doctorado\Primer Artículo\Tercera correción-Traducir y enviar\Artículo definitivo en inglés..docx"
out_path = r"C:\Users\leoga\Desktop\Doctorado 2025\Artículos\Articulos Finales Doctorado\Primer Artículo\Tercera correción-Traducir y enviar\Artículo definitivo en inglés_Emerald_Format.docx"
title_path = r"C:\Users\leoga\Desktop\Doctorado 2025\Artículos\Articulos Finales Doctorado\Primer Artículo\Tercera correción-Traducir y enviar\Title_Page_Emerald.docx"

# First, create Title Page
title_doc = docx.Document()
title_doc.add_heading('Title Page', 0)
title_doc.add_paragraph('Article Title: Does Digital Transformation Undermine Gastronomic Authenticity? The Mediating Role of Perceived Gastronomic Authenticity in Tourist Loyalty within a World Heritage City\n\n[NOTA PARA EL AUTOR: Añade aquí los nombres de los autores, afiliaciones (institución, ciudad, país), biografías (máx. 100 palabras) y la sección de agradecimientos (Acknowledgements) si corresponde.]')
title_doc.save(title_path)

# Now, modify the main document
doc = docx.Document(doc_path)

structured_abstract = """Purpose: The study investigates the mediating role of perceived gastronomic authenticity in the relationship between digital transformation—encompassing both urban and restaurant dimensions—and visitor satisfaction, social value, and destination recommendations.
Design/methodology/approach: Utilizing a sample of 331 visitors to a World Heritage City, the research employed partial least squares structural equation modeling (PLS-SEM) to analyze a reflective model.
Findings: The results reveal that gastronomic authenticity initiates a causal sequence that significantly influences social value, satisfaction, and tourist loyalty, establishing it as a central component of the experience in heritage destinations.
Originality: This work represents the first empirical examination that simultaneously addresses both dimensions of digitalization as determinants of perceived gastronomic authenticity within the context of Andalusian culinary heritage.
Practical implications: Offering valuable insights for destination managers and enhancing the positioning of gastronomic identity in the global tourism market."""

for p in doc.paragraphs:
    # Remove title if it's there to anonymize
    if "Does Digital Transformation Undermine" in p.text:
        p.text = "Article classification: Research Paper\n\nStructured Abstract"
        p.runs[0].bold = True
    elif "World Heritage City" in p.text and len(p.text) < 30:
        p.text = ""
    elif p.text.strip() == "Abstract":
        p.text = ""
    elif p.text.startswith("The study investigates the mediating role"):
        p.text = structured_abstract
    # Headings replacement
    elif p.text.startswith("1. Introduction"):
        p.text = "Introduction"
        p.runs[0].bold = True
    elif p.text.startswith("2. Conceptual framework"):
        p.text = "Conceptual framework"
        p.runs[0].bold = True
    elif p.text.startswith("2.1. The digitalisation"):
        p.text = "The digitalisation of the restaurant and the city and its relationship with perceived gastronomic authenticity"
        p.runs[0].italic = True
    elif p.text.startswith("2.2. Gastronomic authenticity and"):
        p.text = "Gastronomic authenticity and its relationship with social value"
        p.runs[0].italic = True
    elif p.text.startswith("2.3. Gastronomic authenticity perceived"):
        p.text = "Gastronomic authenticity perceived as a determinant of tourist satisfaction and loyalty"
        p.runs[0].italic = True
    elif p.text.startswith("2.4. Tourist satisfaction"):
        p.text = "Tourist satisfaction as a determinant of loyalty to the gastronomic destination"
        p.runs[0].italic = True
    elif p.text.startswith("3 Methodology"):
        p.text = "Methodology"
        p.runs[0].bold = True
    elif p.text.startswith("3.1 Research context"):
        p.text = "Research context"
        p.runs[0].italic = True
    elif p.text.startswith("3.2 Methods of analysis"):
        p.text = "Methods of analysis and participants"
        p.runs[0].italic = True
    elif p.text.startswith("3.3. Measurement instrument"):
        p.text = "Measurement instrument and questionnaire design"
        p.runs[0].italic = True
    elif p.text.startswith("4. Research results"):
        p.text = "Research results"
        p.runs[0].bold = True
    elif p.text.startswith("4.1. Characteristics of the tourist"):
        p.text = "Characteristics of the tourist"
        p.runs[0].italic = True
    elif p.text.startswith("4.1. Results of the proposed model"):
        p.text = "Results of the proposed model"
        p.runs[0].italic = True
    elif p.text.startswith("5. Discussion"):
        p.text = "Discussion"
        p.runs[0].bold = True
    elif p.text.startswith("5.1. Implications for research"):
        p.text = "Implications for research in heritage gastronomic tourism"
        p.runs[0].italic = True
    elif p.text.startswith("5.2. Implications for the management"):
        p.text = "Implications for the management of heritage tourist destinations"
        p.runs[0].italic = True
    elif p.text.startswith("6. Conclusions"):
        p.text = "Conclusions"
        p.runs[0].bold = True
    elif p.text.startswith("REFERENCES"):
        p.text = "References"
        p.runs[0].bold = True

    # Tables in text. Using string replace directly on runs could be split, but most likely Table and Number are together.
    # To be safe, we replace it on the paragraph text and lose formatting only for those few paragraphs where tables are mentioned.
    # Wait, losing formatting on text paragraphs is bad. Let's try doing it on runs.
    for run in p.runs:
        if "Table 1" in run.text:
            run.text = run.text.replace("Table 1", "Table I")
        if "Table 2" in run.text:
            run.text = run.text.replace("Table 2", "Table II")
        if "Table 3" in run.text:
            run.text = run.text.replace("Table 3", "Table III")
        if "Table 4" in run.text:
            run.text = run.text.replace("Table 4", "Table IV")
        if "Table 5" in run.text:
            run.text = run.text.replace("Table 5", "Table V")

# Also, update table captions which are their own paragraphs.
for p in doc.paragraphs:
    if p.text.startswith("Table I."):
        p.text = "[Insert Table I here]\n" + p.text
    elif p.text.startswith("Table II."):
        p.text = "[Insert Table II here]\n" + p.text
    elif p.text.startswith("Table III."):
        p.text = "[Insert Table III here]\n" + p.text
    elif p.text.startswith("Table IV."):
        p.text = "[Insert Table IV here]\n" + p.text
    elif p.text.startswith("Table V."):
        p.text = "[Insert Table V here]\n" + p.text
    elif p.text.startswith("Figure 1."):
        p.text = "[Insert Figure 1 here]\n" + p.text

doc.save(out_path)
print("Formatting complete")
